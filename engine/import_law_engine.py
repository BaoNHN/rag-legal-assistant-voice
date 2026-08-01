# import_law_engine.py
# Background worker: extract text (or OCR) a PDF and add to ChromaDB.
# Auto-detects selectable PDFs and skips OCR when direct text extraction suffices.

import os
import threading
import warnings
import unicodedata, re as _re
import cv2
import numpy as np
import torch

from pdf2image import convert_from_path
from pypdf import PdfReader
from langchain.schema import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from vietocr.tool.predictor import Predictor
from vietocr.tool.config import Cfg
from PIL import ImageEnhance

from database.database import upsert_import_chat

warnings.filterwarnings("ignore", message=".*enable_nested_tensor.*", category=UserWarning)
warnings.filterwarnings("ignore", message=".*batch_first.*", category=UserWarning)

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))  # engine/ → root
DB_PATH = os.path.join(BASE_DIR, "chroma_db")
POPPLER_PATH = os.path.join(BASE_DIR, "poppler", "Library", "bin")
DEVICE_CFG = os.path.join(BASE_DIR, ".device_config")
INSERT_BATCH = 32

# ── Job status registry (job_id → status dict) ──────────────────────────────
_jobs: dict = {}
_jobs_lock = threading.Lock()


def get_job(job_id: str) -> dict:
    with _jobs_lock:
        return _jobs.get(job_id, {})


def _set_job(job_id: str, **kwargs):
    with _jobs_lock:
        if job_id not in _jobs:
            _jobs[job_id] = {}
        _jobs[job_id].update(kwargs)


# ── Device detection ─────────────────────────────────────────────────────────
def _detect_device() -> str:
    if os.path.exists(DEVICE_CFG):
        with open(DEVICE_CFG) as f:
            for line in f.read().splitlines():
                if line.startswith("DEVICE="):
                    val = line.split("=", 1)[1].strip().lower()
                    if val == "cuda":
                        try:
                            if torch.cuda.is_available():
                                return "cuda"
                        except ImportError:
                            pass
                    return "cpu"
    try:
        return "cuda" if torch.cuda.is_available() else "cpu"
    except ImportError:
        return "cpu"


# ── Image preprocessing ───────────────────────────────────────────────────────
def _preprocess(pil_image):
    img = pil_image.convert('L')
    img = ImageEnhance.Contrast(img).enhance(1.8)
    img = ImageEnhance.Sharpness(img).enhance(2.0)
    return img.convert('RGB')


# ── OpenCV line detection ─────────────────────────────────────────────────────
def _detect_lines(pil_image):
    try:

        arr = np.array(pil_image.convert('RGB'))
        gray = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
        img_h, img_w = gray.shape

        _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)

        TOP_MARGIN = 320
        LEFT_MARGIN = 150
        RIGHT_EDGE = img_w - 150
        MIN_WIDTH = 150
        MAX_LINE_H = 110
        SKIP_H = 400

        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (80, 1))
        dilated = cv2.dilate(binary, kernel, iterations=1)
        contours, _ = cv2.findContours(dilated, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

        raw = []
        for cnt in contours:
            x, y, w, h = cv2.boundingRect(cnt)
            if h >= 10 and w >= MIN_WIDTH and x >= LEFT_MARGIN and (x + w) <= RIGHT_EDGE and y >= TOP_MARGIN:
                raw.append((y, y + h))
        raw.sort()

        final = []
        for y1, y2 in raw:
            bh = y2 - y1
            if bh > SKIP_H:
                continue
            if bh <= MAX_LINE_H:
                final.append((max(0, y1 - 4), min(img_h, y2 + 4)))
            else:
                region = binary[y1:y2, :]
                sub_dil = cv2.dilate(region, kernel, iterations=1)
                sc, _ = cv2.findContours(sub_dil, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
                subs = []
                for c in sc:
                    sx, sy, sw, sh = cv2.boundingRect(c)
                    if sh >= 10 and sw >= MIN_WIDTH:
                        subs.append((y1 + sy, y1 + sy + sh))
                if subs:
                    subs.sort()
                    for sy1, sy2 in subs:
                        if (sy2 - sy1) <= MAX_LINE_H:
                            final.append((max(0, sy1 - 4), min(img_h, sy2 + 4)))
                elif bh <= MAX_LINE_H * 2:
                    final.append((max(0, y1 - 4), min(img_h, y2 + 4)))

        final.sort()
        return final
    except ImportError:
        return None


# ── Line validity filter ──────────────────────────────────────────────────────
def _is_valid_line(text: str) -> bool:
    if not text or len(text) < 3:
        return False
    letters = sum(1 for c in text if unicodedata.category(c).startswith('L'))
    if letters / len(text) < 0.3:
        return False
    words = text.split()
    if len(words) <= 2:
        vowels = set('aăâeêiouôơưAĂÂEÊIOUÔƠƯáàảãạắằẳẵặấầẩẫậéèẻẽẹếềểễệíìỉĩịóòỏõọốồổỗộớờởỡợúùủũụứừửữựýỳỷỹỵ'
                     'ÁÀẢÃẠẮẰẲẴẶẤẦẨẪẬÉÈẺẼẸẾỀỂỄỆÍÌỈĨỊÓÒỎÕỌỐỒỔỖỘỚỜỞỠỢÚÙỦŨỤỨỪỬỮỰÝỲỶỸỴ')
        if not any(c in vowels for c in text):
            return False
    if _re.search(r'\s{5,}[.\s,]{5,}', text):
        return False
    return True


# ── OCR a single page ─────────────────────────────────────────────────────────
def _ocr_page(detector, pil_image) -> str:
    img = _preprocess(pil_image)
    width = img.width
    boxes = _detect_lines(img)
    if not boxes:
        return detector.predict(img)
    lines = []
    for y1, y2 in boxes:
        if (y2 - y1) < 10:
            continue
        crop = img.crop((0, y1, width, y2))
        try:
            t = detector.predict(crop)
            if t and _is_valid_line(t.strip()):
                lines.append(t.strip())
        except Exception:
            pass
    return "\n".join(lines)


# ── Sub-split an over-long single article ─────────────────────────────────────
# A handful of articles (e.g. Điều 74) run to 10,000+ characters — one chunk
# that long makes the embedding represent too much at once and forces the LLM
# to read a huge context. This splits only within a single over-long article,
# never merging content across articles: short/normal articles pass through
# _segment() untouched (single chunk), only chunks > 3000 chars get sub-split
# here into 2000–3000 char pieces with 200–300 char overlap.
_ARTICLE_HEADER_RE = _re.compile(r'^(Điều\s+\d+[a-z]?[.,]\s*[^\n]*)', _re.IGNORECASE)


def _split_long_segment(seg: str, size: int = 2800, overlap: int = 250) -> list:
    if len(seg) <= 3000:
        return [seg]

    header_match = _ARTICLE_HEADER_RE.match(seg)
    header = header_match.group(1).strip() if header_match else ""

    pieces = []
    i, n = 0, len(seg)
    while i < n:
        end = min(i + size, n)
        if end < n:
            # Prefer cutting at a paragraph/sentence boundary near the
            # target size instead of mid-word/mid-sentence.
            boundary = seg.rfind('\n', i + 500, end)
            if boundary == -1:
                boundary = seg.rfind('. ', i + 500, end)
            if boundary != -1:
                end = boundary + 1
        piece = seg[i:end].strip()
        if piece:
            pieces.append(piece)
        if end >= n:
            break
        i = end - overlap

    # run_import() re-derives article_number/article_reference per chunk by
    # regex-matching "Điều N." at the start of its text — without this, only
    # the first piece would carry that header and every later piece would
    # silently fall back to an untagged (uncitable) chunk.
    if header:
        for k in range(1, len(pieces)):
            if not _ARTICLE_HEADER_RE.match(pieces[k]):
                pieces[k] = f"{header} (tiếp theo)\n{pieces[k]}"

    return pieces


# ── Segment full text into articles ──────────────────────────────────────────
def _segment(full_text: str) -> tuple:
    """Split into one chunk per 'Điều X.' (legal article), sub-splitting any
    single article over 3000 chars (see _split_long_segment above).

    Returns (segments, matched_by_article). matched_by_article is False when
    the article-boundary regex couldn't find enough headers and a fixed-size
    fallback chunking was used instead — fallback chunks straddle article
    boundaries and must NOT be tagged with a fabricated article number.
    """
    clean = _re.sub(r'\n{3,}', '\n\n', full_text)
    clean = _re.sub(r'[ \t]+', ' ', clean).strip()
    pattern = r'(?:(?:^|\n)(?=Điều\s+\d+[a-z]?[.,]\s))'
    splits = _re.split(pattern, clean, flags=_re.MULTILINE)
    segs = [s.strip() for s in splits if len(s.strip()) > 50]
    if len(segs) >= 5:
        expanded = []
        for s in segs:
            expanded.extend(_split_long_segment(s))
        return expanded, True

    # fallback: fixed chunks — article boundaries could not be detected.
    # Stays at the coarser 3000/300 sizing (not 2000-3000/200-300) since
    # these chunks straddle unknown article boundaries anyway — there's no
    # single article's worth of content to preserve intact here.
    size, overlap, segs = 3000, 300, []
    i = 0
    while i < len(clean):
        segs.append(clean[i:i + size])
        i += size - overlap
    return segs, False


# ── Direct text extraction (selectable PDFs) ─────────────────────────────────
_MIN_CHARS_PER_PAGE = 150  # avg below this → treat as scanned, use OCR


def _extract_text_pdf(pdf_path: str, total_pages: int):
    """
    Try to extract text directly with pypdf.
    Returns {page_num: text} when the PDF has selectable text,
    or None when it appears to be a scan (too few characters).
    """
    reader = PdfReader(pdf_path)
    pages_text = {}
    for i, page in enumerate(reader.pages, start=1):
        raw = page.extract_text() or ""
        text = _re.sub(r'[ \t]+', ' ', raw)
        text = _re.sub(r'\n{3,}', '\n\n', text).strip()
        pages_text[i] = text

    avg_chars = sum(len(t) for t in pages_text.values()) / max(total_pages, 1)
    return pages_text if avg_chars >= _MIN_CHARS_PER_PAGE else None


def _extract_text_docx(docx_path: str) -> dict:
    """Extract text from a .docx Word document. Returns {1: full_text}."""
    from docx import Document as DocxDoc
    doc = DocxDoc(docx_path)

    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # Also pull text from tables (law docs often have tabular articles)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in lines:
                    lines.append(text)

    return {1: "\n".join(lines)}


# ── Main background job ───────────────────────────────────────────────────────
def run_import(job_id: str, file_path: str, so_ky_hieu: str,
               loai_van_ban: str, nguon_thu_thap: str,
               student_id: int, db_conn_factory, importer: str = "admin1",
               primary_keyword_ids: list = None, secondary_keyword_ids: list = None):
    """
    Full pipeline: auto-detect file type (PDF / DOCX) → extract text or OCR →
    segment → embed → add to ChromaDB (no wipe).
    """
    ext = os.path.splitext(file_path)[1].lower()
    _set_job(job_id, status="running", message="Kiểm tra loại file…")

    try:
        if ext == ".docx":
            # ── DOCX: direct text extraction, no OCR needed ──
            _set_job(job_id, message="File Word (.docx) — đang trích xuất văn bản…")
            all_text_by_page = _extract_text_docx(file_path)
            total_pages = 1
            _set_job(job_id, message=f"Trích xuất DOCX hoàn tất ({len(all_text_by_page[1]):,} ký tự).")
        else:
            # ── PDF: try direct text, fall back to OCR ──
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)

            # ── 1. Try direct text extraction first ──
            all_text_by_page = _extract_text_pdf(file_path, total_pages)

            if all_text_by_page:
                _set_job(job_id, message=f"PDF văn bản số ({total_pages} trang). Trích xuất hoàn tất.")
            else:
                # ── 2. Fall back to VietOCR for scanned PDFs ──
                _set_job(job_id, message="PDF scan — đang khởi động VietOCR…")
                device = _detect_device()
                cfg = Cfg.load_config_from_name('vgg_transformer')
                cfg['device'] = device
                cfg['predictor']['beamsearch'] = True
                detector = Predictor(cfg)
                _set_job(job_id, message=f"VietOCR loaded ({device.upper()}). Bắt đầu OCR…")

                all_text_by_page = {}
                BATCH = 5
                DPI = 250

                for batch_start in range(1, total_pages + 1, BATCH):
                    batch_end = min(batch_start + BATCH - 1, total_pages)
                    _set_job(job_id, message=f"OCR trang {batch_start}–{batch_end}/{total_pages}…")

                    kwargs = dict(dpi=DPI, first_page=batch_start, last_page=batch_end, fmt='jpeg')
                    if os.path.isdir(POPPLER_PATH):
                        kwargs['poppler_path'] = POPPLER_PATH

                    pages = convert_from_path(file_path, **kwargs)
                    for i, pg in enumerate(pages):
                        all_text_by_page[batch_start + i] = _ocr_page(detector, pg)
                    del pages

        # ── 3. Save raw text ──
        raw_txt = os.path.join(BASE_DIR, f"{so_ky_hieu.replace('/', '_')}_{nguon_thu_thap.replace(' ', '_')}.txt")
        with open(raw_txt, 'w', encoding='utf-8') as f:
            for p in sorted(all_text_by_page):
                f.write(f"\n\n=== TRANG {p} ===\n{all_text_by_page[p]}")

        full_text = "\n".join(all_text_by_page[p] for p in sorted(all_text_by_page))

        # ── 4. Segment ──
        _set_job(job_id, message=f"Phân đoạn văn bản pháp luật ({len(full_text):,} ký tự)…")
        segs, matched_by_article = _segment(full_text)

        if not matched_by_article:
            _set_job(job_id, message=(
                "⚠️ Không nhận diện được ranh giới 'Điều X.' trong văn bản — "
                "đang dùng chế độ cắt đoạn dự phòng (3000 ký tự/đoạn). "
                "Trích dẫn theo số Điều có thể không chính xác cho văn bản này."
            ))

        # ── 5. Build documents ──
        from engine.rag_engine import detect_entity_type

        docs = []
        for i, seg in enumerate(segs):
            m = _re.match(r'Điều\s+(\d+[a-z]?)[.,\s]', seg)
            lines = [l.strip() for l in seg.split('\n') if l.strip()]
            meta = {
                "so_ky_hieu": so_ky_hieu,
                "loai_van_ban": loai_van_ban,
                "nguon_thu_thap": nguon_thu_thap,
                "char_count": len(seg),
                "segment_index": i,
                "import_source": "law",
                "importer": importer,
            }
            if m:
                # Real article boundary — safe to tag with its true number.
                art_num = m.group(1)
                meta["article_number"] = art_num
                meta["article_reference"] = f"Điều {art_num}"
                title = lines[0][:120] if lines else f"Điều {art_num}"
            else:
                # Fallback chunk with no detected header — do NOT fabricate an
                # article number (segment index != real article number).
                title = lines[0][:120] if lines else f"Đoạn {i + 1}"
            meta["title"] = title

            # retrieval_keywords: without this, raw law-text chunks score
            # only on flat content keyword-overlap in select_best_doc(),
            # while curated dataset chunks (KB_Articles_Updated etc.) carry
            # a dense manually-written retrieval_keywords field worth 3x per
            # word plus a +10 exact-phrase bonus — official law chunks lost
            # that comparison every time despite being the higher-tier
            # source (see nhận xét 25/7 finding). Derive a topic phrase from
            # the article's own heading (stripping the "Điều N." prefix) so
            # it gets the same scoring treatment.
            topic_phrase = _re.sub(r'^Điều\s+\d+[a-z]?[.,]\s*', '', title).strip()
            if topic_phrase:
                meta["retrieval_keywords"] = topic_phrase

            # Stamp entity_type upfront so infer_doc_entity() hits the cheap
            # explicit-field path at query time instead of re-scanning seg
            # text on every retrieval (see rag_engine.backfill_entity_type_tags
            # docstring — matters once the corpus grows to tens of thousands
            # of chunks).
            entity_type = detect_entity_type(f"{title} {seg[:500]}")
            if entity_type:
                meta["entity_type"] = entity_type

            docs.append(Document(page_content=seg, metadata=meta))

        # ── 6. Add to ChromaDB (no wipe, skip existing so_ky_hieu) ──
        _set_job(job_id, message=f"Tải {len(docs)} đoạn lên ChromaDB…")

        embedding = HuggingFaceEmbeddings(
            model_name="BAAI/bge-m3",
            model_kwargs={"device": "cpu"},
            encode_kwargs={"normalize_embeddings": True},
        )
        vs = Chroma(persist_directory=DB_PATH, embedding_function=embedding)

        # Skip if already exists
        existing = vs.get(include=["metadatas"])
        existing_ids = {m.get("so_ky_hieu", "").strip() for m in existing["metadatas"]}

        new_docs = [d for d in docs if d.metadata["so_ky_hieu"] not in existing_ids]
        skipped = len(docs) - len(new_docs)

        if new_docs:
            for i in range(0, len(new_docs), INSERT_BATCH):
                chunk = new_docs[i:i + INSERT_BATCH]
                vs.add_documents(chunk)
                _set_job(job_id, message=f"Indexed {min(i + INSERT_BATCH, len(new_docs))}/{len(new_docs)}…")

        result_msg = (
            f"✅ Hoàn tất! Đã thêm {len(new_docs)} đoạn vào ChromaDB "
            f"(bỏ qua {skipped} đoạn trùng lặp)."
        )
        if not matched_by_article:
            result_msg += (
                "\n⚠️ Lưu ý: không tách được theo 'Điều X.' — đã dùng cắt đoạn dự phòng, "
                "trích dẫn số Điều có thể không chính xác cho văn bản này."
            )

        # Refresh the citation-source whitelist so the new so_ky_hieu becomes
        # citable immediately (see engine.rag_engine.refresh_citation_sources).
        from engine.rag_engine import refresh_citation_sources
        refresh_citation_sources()

        # Tag this so_ky_hieu with its admin/teacher-selected keywords (see
        # database.database.keyword / source_keyword) — written unconditionally,
        # even when every chunk was skipped as a duplicate, so re-submitting an
        # already-imported so_ky_hieu can still be used just to update its tags.
        from database.database import set_source_keywords
        set_source_keywords("law", so_ky_hieu, primary_keyword_ids or [], secondary_keyword_ids or [])

        _set_job(job_id, status="done", message=result_msg)

    except Exception as e:
        import traceback
        traceback.print_exc()
        _set_job(job_id, status="failed", message=f"❌ Lỗi: {e}")

    finally:
        # ── 7. Create/update 'Import new law' chat ──
        try:
            msg = get_job(job_id).get("message", "Xử lý hoàn tất.")
            upsert_import_chat(student_id, msg)
        except Exception:
            pass
        # Cleanup temp file
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception:
            pass