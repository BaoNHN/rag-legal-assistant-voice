# build_db_doc.py
# Import a single PDF or DOCX (selectable text) into ChromaDB.
# No OCR needed — uses direct text extraction.
#
# USAGE:
#   conda activate rag_env
#   cd D:\hoc\project\rag-legal-assistant-master
#
#   python database/build_db_doc.py "67_VBHN-VPQH_671127 (1).docx" --so-ky-hieu "67/VBHN-VPQH" --loai "Van ban hop nhat" --nguon "vbpl.vn"
#   python database/build_db_doc.py "2025_1107 + 1108_67-VBHN-VPQH.pdf" --so-ky-hieu "67/VBHN-VPQH" --loai "Van ban hop nhat" --nguon "chinhphu.vn"

import os
import re
import sys
import argparse

from langchain.schema import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma

BASE_DIR     = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DB_PATH      = os.path.join(BASE_DIR, "chroma_db")
INSERT_BATCH = 32
MIN_CHARS_PAGE = 150   # below this avg → warn that file may need OCR


# ── Text extractors ───────────────────────────────────────────────────────────
def extract_pdf(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    pages  = []
    for i, page in enumerate(reader.pages, 1):
        raw  = page.extract_text() or ""
        text = re.sub(r'[ \t]+', ' ', raw)
        text = re.sub(r'\n{3,}', '\n\n', text).strip()
        pages.append(text)
        if i % 20 == 0:
            print(f"  Read {i}/{len(reader.pages)} pages...")

    avg = sum(len(p) for p in pages) / max(len(pages), 1)
    if avg < MIN_CHARS_PAGE:
        print(f"[!] Low avg chars/page ({avg:.0f}). File may be scanned.")
        print("    Consider using build_db_from_pdf.py (VietOCR) instead.")
    return "\n\n".join(pages)


def extract_docx(path: str) -> str:
    from docx import Document as DocxDoc
    doc   = DocxDoc(path)
    lines = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in lines:
                    lines.append(t)
    return "\n".join(lines)


# ── Segmenter ─────────────────────────────────────────────────────────────────
def segment(full_text: str) -> tuple:
    """Split into one chunk per 'Điều X.' (legal article).

    Returns (segments, matched_by_article). matched_by_article is False when
    the article-boundary regex couldn't find enough headers and a fixed-size
    fallback chunking was used instead — fallback chunks straddle article
    boundaries and must NOT be tagged with a fabricated article number.
    """
    clean = re.sub(r'\n{3,}', '\n\n', full_text)
    clean = re.sub(r'[ \t]+', ' ', clean).strip()

    # Split on "Điều X." boundaries (real Vietnamese "Đ", not ASCII "D")
    pattern = r'(?:(?:^|\n)(?=Điều\s+\d+[a-z]?[.,]\s))'
    parts   = re.split(pattern, clean, flags=re.MULTILINE)
    segs    = [s.strip() for s in parts if len(s.strip()) > 50]

    if len(segs) >= 5:
        return segs, True

    # Fallback: fixed-size chunks — article boundaries could not be detected
    size, overlap, segs = 3000, 300, []
    i = 0
    while i < len(clean):
        segs.append(clean[i:i + size])
        i += size - overlap

    return segs, False


# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(
        description="Import a PDF or DOCX (with selectable text) into ChromaDB"
    )
    parser.add_argument("file", help="File name in project root, or full path")
    parser.add_argument("--so-ky-hieu", default="",
                        help='Legal document ID, e.g. "67/VBHN-VPQH"')
    parser.add_argument("--loai", default="Van ban phap luat",
                        help='Document type, e.g. "Van ban hop nhat"')
    parser.add_argument("--nguon", default="",
                        help='Source, e.g. "vbpl.vn" or "chinhphu.vn"')
    parser.add_argument("--force", action="store_true",
                        help="Re-import even if so_ky_hieu already exists in DB")
    args = parser.parse_args()

    # Resolve path
    file_path = args.file if os.path.isabs(args.file) \
                else os.path.join(BASE_DIR, args.file)

    if not os.path.exists(file_path):
        print(f"[ERROR] File not found: {file_path}")
        sys.exit(1)

    ext         = os.path.splitext(file_path)[1].lower()
    so_ky_hieu  = args.so_ky_hieu.strip()
    loai        = args.loai.strip()
    nguon       = args.nguon.strip() or os.path.basename(file_path)

    if not so_ky_hieu:
        # Derive from filename if not provided
        so_ky_hieu = re.sub(r'[^A-Za-z0-9/]', '', os.path.splitext(os.path.basename(file_path))[0])
        print(f"[!] --so-ky-hieu not set, using: {so_ky_hieu}")

    print(f"\nFile      : {os.path.basename(file_path)}")
    print(f"So ky hieu: {so_ky_hieu}")
    print(f"Loai      : {loai}")
    print(f"Nguon     : {nguon}")
    print()

    # ── Extract text ──
    print("Extracting text...")
    if ext == ".docx":
        full_text = extract_docx(file_path)
    elif ext == ".pdf":
        full_text = extract_pdf(file_path)
    else:
        print(f"[ERROR] Unsupported format: {ext}. Use .pdf or .docx")
        sys.exit(1)

    print(f"  Extracted {len(full_text):,} characters")

    # ── Segment ──
    print("Segmenting into articles...")
    segs, matched_by_article = segment(full_text)
    print(f"  Found {len(segs)} segments")
    if not matched_by_article:
        print("[!] Could not detect 'Điều X.' boundaries — using fixed-size fallback chunking.")
        print("    Article-number citations will be inaccurate for this document.")

    # ── Build documents ──
    docs = []
    for i, seg in enumerate(segs):
        m     = re.match(r'Điều\s+(\d+[a-z]?)[.,\s]', seg)
        lines = [l.strip() for l in seg.split('\n') if l.strip()]
        meta = {
            "so_ky_hieu":     so_ky_hieu,
            "loai_van_ban":   loai,
            "nguon_thu_thap": nguon,
            "char_count":     len(seg),
            "segment_index":  i,
        }
        if m:
            # Real article boundary — safe to tag with its true number.
            art_num = m.group(1)
            meta["article_number"] = art_num
            meta["article_reference"] = f"Điều {art_num}"
            meta["title"] = lines[0][:120] if lines else f"Điều {art_num}"
        else:
            # Fallback chunk with no detected header — do NOT fabricate an
            # article number (segment index != real article number).
            meta["title"] = lines[0][:120] if lines else f"Đoạn {i + 1}"

        docs.append(Document(page_content=seg, metadata=meta))

    # ── Load ChromaDB ──
    print("\nLoading embedding model (BAAI/bge-small-en-v1.5)...")
    embedding = HuggingFaceEmbeddings(model_name="BAAI/bge-small-en-v1.5")
    vs        = Chroma(persist_directory=DB_PATH, embedding_function=embedding)

    # ── Dedup check ──
    existing    = vs.get(include=["metadatas"])
    existing_ids = {m.get("so_ky_hieu", "").strip() for m in existing["metadatas"]}

    if so_ky_hieu in existing_ids and not args.force:
        print(f"\n[SKIP] '{so_ky_hieu}' already exists in ChromaDB.")
        print("       Use --force to re-import and overwrite.")
        print(f"       Total in DB: {vs._collection.count()}")
        sys.exit(0)

    new_docs = [d for d in docs if d.metadata["so_ky_hieu"] not in existing_ids] \
               if not args.force else docs

    skipped = len(docs) - len(new_docs)
    if skipped:
        print(f"  Skipping {skipped} segments (already in DB)")

    # ── Index ──
    if not new_docs:
        print("Nothing new to add.")
    else:
        print(f"\nIndexing {len(new_docs)} segments into ChromaDB...")
        for i in range(0, len(new_docs), INSERT_BATCH):
            chunk = new_docs[i:i + INSERT_BATCH]
            vs.add_documents(chunk)
            done = min(i + INSERT_BATCH, len(new_docs))
            print(f"  Indexed {done}/{len(new_docs)}")

    total = vs._collection.count()
    print(f"\n[DONE]")
    print(f"  Inserted : {len(new_docs)} segments")
    print(f"  Total DB : {total}")


if __name__ == "__main__":
    main()
