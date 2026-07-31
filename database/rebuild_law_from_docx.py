# rebuild_law_from_docx.py
# Fixes broken chunking of the Enterprise Law consolidated text (67/VBHN-VPQH):
# the previous import fell back to fixed 3000-char sliding-window chunks
# (article_number metadata == segment_index+1, i.e. meaningless), splitting
# article content across chunk boundaries and causing wrong answers/citations.
#
# This script:
#   1. Removes the old fallback-chunked law documents (so_ky_hieu
#      '67/VBHN-VPQH' and '67/VBHN-VPQH-2025', identified by presence of
#      segment_index metadata) — leaves Q&A / KB_Articles dataset docs untouched.
#   2. Re-extracts text from the clean DOCX source and re-segments it per
#      "Điều N." boundary (verified: 221 segments, 1 preamble + 220 articles).
#   3. Re-adds correctly chunked documents with accurate article_number /
#      article_reference metadata.
#
# USAGE:
#   conda activate rag_env
#   cd D:\hoc\project\rag-legal-assistant-master
#   python database/rebuild_law_from_docx.py

import os
import re
from docx import Document as DocxDoc
from langchain.schema import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCX_PATH = os.path.join(BASE_DIR, "67_VBHN-VPQH_671127 (1).docx")
DB_PATH = os.path.join(BASE_DIR, "chroma_db")
INSERT_BATCH = 32

SO_KY_HIEU = "67/VBHN-VPQH"
LOAI_VAN_BAN = "Văn bản hợp nhất"
NGUON_THU_THAP = "vbpl.vn"


def extract_docx_text(path: str) -> str:
    doc = DocxDoc(path)
    lines = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in lines:
                    lines.append(t)
    return "\n".join(lines)


def segment_by_article(text: str) -> list:
    clean = re.sub(r'\n{3,}', '\n\n', text)
    clean = re.sub(r'[ \t]+', ' ', clean).strip()
    pattern = r'(?:(?:^|\n)(?=Điều\s+\d+[a-z]?[.,]\s))'
    splits = re.split(pattern, clean, flags=re.MULTILINE)
    return [s.strip() for s in splits if len(s.strip()) > 50]


def main():
    print(f"Loading DOCX: {DOCX_PATH}")
    text = extract_docx_text(DOCX_PATH)
    print(f"  Extracted {len(text):,} characters")

    segs = segment_by_article(text)
    print(f"  Segmented into {len(segs)} pieces")

    docs = []
    no_match = 0
    for i, seg in enumerate(segs):
        m = re.match(r'Điều\s+(\d+[a-z]?)[.,\s]', seg)
        if m:
            art_num = m.group(1)
        else:
            art_num = None
            no_match += 1

        lines = [l.strip() for l in seg.split('\n') if l.strip()]
        title = lines[0][:120] if lines else (f"Điều {art_num}" if art_num else f"Đoạn {i+1}")

        meta = {
            "so_ky_hieu": SO_KY_HIEU,
            "loai_van_ban": LOAI_VAN_BAN,
            "nguon_thu_thap": NGUON_THU_THAP,
            "title": title,
            "char_count": len(seg),
            "segment_index": i,
        }
        if art_num:
            meta["article_number"] = art_num
            meta["article_reference"] = f"Điều {art_num}"

        docs.append(Document(page_content=seg, metadata=meta))

    print(f"  Built {len(docs)} documents ({no_match} without a matched article number, e.g. preamble)")

    print("\nLoading embedding model (BAAI/bge-small-en-v1.5)...")
    embedding = HuggingFaceEmbeddings(model_name="BAAI/bge-small-en-v1.5")
    vs = Chroma(persist_directory=DB_PATH, embedding_function=embedding)

    # ── Remove old fallback-chunked law docs ──
    print("\nFinding old fallback-chunked law documents to remove...")
    existing = vs.get(include=["metadatas"])
    old_ids = [
        _id for _id, meta in zip(existing["ids"], existing["metadatas"])
        if meta.get("so_ky_hieu") in ("67/VBHN-VPQH", "67/VBHN-VPQH-2025")
        and "segment_index" in meta
    ]
    print(f"  Found {len(old_ids)} old documents to delete")
    if old_ids:
        for i in range(0, len(old_ids), 200):
            vs.delete(ids=old_ids[i:i + 200])
        print(f"  Deleted {len(old_ids)} old documents")

    # ── Add new correctly-chunked docs ──
    print(f"\nAdding {len(docs)} new documents to ChromaDB...")
    for i in range(0, len(docs), INSERT_BATCH):
        chunk = docs[i:i + INSERT_BATCH]
        vs.add_documents(chunk)
        print(f"  Indexed {min(i + INSERT_BATCH, len(docs))}/{len(docs)}")

    final = vs.get(include=["metadatas"])
    print(f"\n✅ DONE. Total documents in ChromaDB now: {len(final['metadatas'])}")


if __name__ == "__main__":
    main()
